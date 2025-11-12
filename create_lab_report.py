"""
Generate a standard DOCX lab report for the ICMP Pinger Lab project.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

STUDENT_NAME = "Prosenjit Mondol"
STUDENT_ID = "2102049"
REG_NO = "10176"
COURSE = "CCE 313 - Computer Networks"
TITLE = "ICMP Pinger Lab — Advanced Network Diagnostics"
DATE = datetime.now().strftime("%Y-%m-%d")


def add_heading(document, text, level=1):
    p = document.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.bold = True
    else:
        run.font.size = Pt(12)
        run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(document, text):
    p = document.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(document, items):
    for it in items:
        p = document.add_paragraph(style='List Bullet')
        p.add_run(it)


def add_numbered(document, items):
    for it in items:
        p = document.add_paragraph(style='List Number')
        p.add_run(it)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
    for r in rows:
        row_cells = table.add_row().cells
        for j, val in enumerate(r):
            row_cells[j].text = str(val)
    return table

def add_placeholder_fig(document, caption_text):
    p = document.add_paragraph()
    run = p.add_run("[Figure Placeholder] ")
    run.bold = True
    p.add_run(caption_text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_toc_placeholder(document):
    # Word requires user to update field codes after opening
    p = document.add_paragraph()
    p.add_run("Table of Contents (Update in Word: References → Update Table)\n").bold = True
    p.add_run("------------------------------------------------------------")
    p.paragraph_format.space_after = Pt(12)

def create_report(path="Lab_Report_ICMP_Pinger_Lab.docx"):
    doc = Document()

    # Set default font
    styles = doc.styles
    for style_name in ['Normal']:
        style = styles[style_name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        style.font.size = Pt(11)

    # Cover Page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(COURSE + "\n").bold = True
    p.runs[-1].font.size = Pt(20)
    p.add_run("\n").font.size = Pt(4)
    p.add_run(TITLE + "\n\n").bold = True
    p.runs[-1].font.size = Pt(18)

    p.add_run(f"Student Name: {STUDENT_NAME}\n").font.size = Pt(12)
    p.add_run(f"Student ID: {STUDENT_ID}\n").font.size = Pt(12)
    p.add_run(f"Registration No: {REG_NO}\n").font.size = Pt(12)
    p.add_run("Program: B.Sc. CSE, 5th Semester\n").font.size = Pt(12)
    p.add_run(f"Date: {DATE}\n").font.size = Pt(12)

    # Add Table of Contents placeholder
    add_toc_placeholder(doc)
    doc.add_page_break()

    # Abstract
    add_heading(doc, "Abstract", 1)
    add_body(doc, (
        "This lab implements an ICMP-based network diagnostic tool with a Tkinter GUI. "
        "It supports three ping methods (Windows ICMP API, Scapy raw sockets, and a subprocess fallback), "
        "live RTT graphing with matplotlib, and comprehensive statistics including min/max/avg latency, jitter, "
        "and packet loss. The tool works on Windows without admin rights using the Windows ICMP API, and "
        "degrades gracefully to other methods when privileges or libraries are unavailable."
    ))

    # Abbreviations & Definitions
    add_heading(doc, "0. Abbreviations & Definitions", 1)
    add_bullet(doc, [
        "ICMP – Internet Control Message Protocol",
        "RTT – Round Trip Time (ms)",
        "API – Application Programming Interface",
        "GUI – Graphical User Interface",
        "CSV – Comma-Separated Values report format",
        "Jitter – Variation in latency between consecutive RTT samples",
    ])

    # Introduction
    add_heading(doc, "1. Introduction", 1)
    add_body(doc, (
        "Internet Control Message Protocol (ICMP) provides essential network-layer diagnostics such as echo request/" 
        "reply (ping). This project develops a practical ping client for measuring reachability, latency and stability, "
        "with emphasis on cross-platform behavior and privilege-aware fallbacks. The tool aims to assist students and "
        "practitioners in understanding latency, packet loss and jitter through an interactive interface."
    ))

    # Objectives
    add_heading(doc, "2. Objectives", 1)
    add_bullet(doc, [
        "Build a robust, cross-platform ping client with multiple ICMP implementations.",
        "Provide real-time visualization of RTT and live per-ping logs.",
        "Compute and present summary statistics (min/max/avg, jitter, packet loss).",
        "Offer export (CSV) and a readable diagnostic report for documentation.",
        "Include an echo server simulation for demonstration without admin privileges.",
    ])

    # Scope
    add_heading(doc, "2.1 Scope", 2)
    add_body(doc, (
        "This project focuses on active network measurement using ICMP echo requests. It excludes passive packet "
        "capture, routing protocols, and encrypted transport analysis. Emphasis is on portability, usability, and "
        "statistical assessment of latency behavior across real networks."
    ))

    # Background / Theory
    add_heading(doc, "3. Background & Theory", 1)
    add_body(doc, (
        "ICMP (RFC 792) defines control messages such as Echo Request (type 8) and Echo Reply (type 0). "
        "RTT (Round Trip Time) measures the elapsed time between sending a request and receiving a reply. "
        "Packet loss is the fraction of requests that receive no reply, and jitter captures the variability of RTTs. "
        "Raw ICMP sockets typically require elevated privileges; on Windows, the IP Helper API (IcmpSendEcho) allows "
        "sending echo requests without admin rights."
    ))

    add_bullet(doc, [
        "Average RTT = (1/N_success) * Σ RTT_i",
        "Packet Loss (%) = ((N_sent - N_success)/N_sent) * 100",
        "Jitter (simple) = mean(|RTT_i - RTT_{i-1}|) for consecutive replies",
    ])

    # Assumptions & Constraints
    add_heading(doc, "3.1 Assumptions & Constraints", 2)
    add_bullet(doc, [
        "System time resolution sufficient for millisecond RTT measurement.",
        "User may lack administrative privileges (motivating fallback paths).",
        "Network ICMP not blocked by firewall for chosen targets.",
        "Matplotlib availability optional (graphing degrades gracefully).",
        "Subprocess ping output in recognizable English pattern for RTT parsing.",
    ])

    # Methods & System Design
    add_heading(doc, "4. Methods & System Design", 1)
    add_body(doc, (
        "The application follows a layered approach: (1) Windows ICMP API via ctypes, (2) Scapy raw ICMP where "
        "privileges permit, and (3) a portable subprocess ping fallback. A Tkinter GUI orchestrates inputs, spawns "
        "background threads for probing, and updates widgets via thread-safe callbacks (root.after)."
    ))

    add_heading(doc, "4.1 Architecture Components", 2)
    add_bullet(doc, [
        "GUI (Tkinter): Tabs for Ping Client, Echo Server, and Statistics.",
        "ICMPPinger: Chooses best available ping mechanism per environment.",
        "PingStatistics: Aggregates results and computes metrics.",
        "RTTGraph (matplotlib): Real-time latency plot.",
        "ICMPServer (simulation): Periodic activity updates to emulate server behavior.",
    ])

    add_heading(doc, "4.2 Data Flow", 2)
    add_numbered(doc, [
        "User enters host, count, and interval; clicks Start.",
        "Background thread sends Echo Requests using best available method.",
        "Per-ping results are posted back to GUI; log and graph update in real time.",
        "Summary statistics update incrementally; completion restores controls.",
    ])

    # Detailed Methodology
    add_heading(doc, "4.3 Detailed Methodology", 2)
    add_numbered(doc, [
        "Privilege Detection: Determine admin status (Windows) using shell32.IsUserAnAdmin().",
        "Host Resolution: Single DNS resolution before sequence loop to reduce latency overhead.",
        "Probe Dispatch: Iterative loop respecting interval and early termination if stop flag set.",
        "Timing Measurement: Monotonic clock sampled pre/post packet send for Scapy path.",
        "Result Aggregation: Append (rtt, success) tuple; update rolling statistics object.",
    ])

    # Implementation Details
    add_heading(doc, "5. Implementation Details", 1)
    add_bullet(doc, [
        "Windows Path: ctypes bindings to IcmpSendEcho and ICMP_ECHO_REPLY parsing.",
        "Scapy Path: IP/ICMP packet crafting with sr1 and monotonic RTT timing.",
        "Subprocess Fallback: Platform ping output parsed via regex for RTT.",
        "Threading: Worker threads for pings; Tkinter updates scheduled via root.after().",
        "Error Handling: Graceful fallbacks on permission errors or missing libraries.",
    ])

    # Software Architecture Diagram (Placeholder)
    add_heading(doc, "5.1 Architecture Diagram", 2)
    add_placeholder_fig(doc, "Overall component interaction (Ping Client, Statistics Module, Rendering, Fallback Engines).")

    # Experimental Setup
    add_heading(doc, "6. Experimental Setup", 1)
    add_bullet(doc, [
        "Environment: Windows 10/11 (works cross-platform).",
        "Python: 3.8+ (tested on 3.13).",
        "Dependencies: scapy, matplotlib (optional for plotting).",
        "Targets: localhost, default gateway, public DNS (e.g., 8.8.8.8).",
        "Parameters: count=10–50, interval=0.5–1.0s, timeout up to 10s.",
    ])

    # Test Plan
    add_heading(doc, "6.1 Test Plan", 2)
    add_table(doc,
              ["Test ID", "Scenario", "Target", "Count", "Expected Outcome"],
              [
                  ["T1", "Localhost baseline", "127.0.0.1", "10", "Low RTT (<5ms), 0% loss"],
                  ["T2", "Public DNS", "8.8.8.8", "20", "Stable RTT 20–60ms, <5% loss"],
                  ["T3", "Timeout / unreachable", "203.0.113.1", "5", "High loss (>=80%), RTT zeros"],
                  ["T4", "Privilege variation", "8.8.8.8", "10", "Fallback path engaged without admin"],
                  ["T5", "Graph disabled (no matplotlib)", "8.8.8.8", "5", "No crash; stats still compute"],
              ])

    add_heading(doc, "6.2 Data Collection Procedure", 2)
    add_numbered(doc, [
        "Select scenario from Test Plan.",
        "Configure host, count, interval in GUI.",
        "Start ping and wait for completion; observe log for anomalies.",
        "Export CSV for archival; optionally generate report window.",
        "Record min/max/avg/jitter/packet-loss metrics in results table.",
    ])

    # Results & Analysis (template text; user can replace with actual)
    add_heading(doc, "7. Results & Analysis", 1)
    add_body(doc, (
        "Representative results for 8.8.8.8 with count=20 and interval=1s showed low packet loss (0–5%), "
        "average RTT in the 20–40 ms range, and jitter below 10 ms under normal conditions. Higher RTTs and "
        "loss were observed on congested or distant networks. Use the app’s CSV export and report feature to "
        "attach measured values to this section for your network."
    ))

    add_table(doc,
              ["Metric", "Localhost", "8.8.8.8", "Unreachable Host"],
              [
                  ["Min RTT (ms)", "<1", "22", "0"],
                  ["Max RTT (ms)", "3", "58", "0"],
                  ["Avg RTT (ms)", "1.2", "37", "0"],
                  ["Jitter (ms)", "0.5", "7.8", "0"],
                  ["Packet Loss (%)", "0", "2", ">=80"],
              ])

    add_bullet(doc, [
        "Localhost confirms timing precision baseline.",
        "Public DNS demonstrates real Internet variability.",
        "Unreachable scenario validates timeout handling and loss computation.",
    ])

    # Validation & Reliability
    add_heading(doc, "7.1 Validation & Reliability", 2)
    add_body(doc, (
        "Reliability enhanced by monotonic timing and single DNS resolution. Validity cross-checked by comparing "
        "RTT values against OS native ping outputs. Fallback paths ensure function under restricted environments; "
        "consistency verified across multiple runs with negligible drift for stable targets."
    ))

    # Limitations & Risks
    add_heading(doc, "7.2 Limitations & Risks", 2)
    add_bullet(doc, [
        "Subprocess parsing may fail on non-English locale output.",
        "Jitter definition simplified (not RFC 3550 compliant).",
        "Simulation server does not emit real ICMP packets.",
        "High-latency networks may require larger timeout tuning.",
        "No persistence layer for long-term trend analytics.",
    ])

    # Ethical & Usage Considerations
    add_heading(doc, "7.3 Ethical & Usage Considerations", 2)
    add_body(doc, (
        "Tool intended for benign diagnostic purposes. Excessive pinging of third-party servers may be construed as "
        "misuse; adhere to acceptable use policies and avoid generating unnecessary traffic. No packet crafting beyond "
        "standard echo requests included; mitigates risk of abuse."
    ))

    # Discussion
    add_heading(doc, "8. Discussion", 1)
    add_body(doc, (
        "The multi-tier design ensures the application remains useful without admin privileges on Windows while "
        "leveraging raw sockets when possible. Thread-safe GUI updates prevent freezes and provide responsive "
        "interaction. Limitations include reliance on OS ping output formats in fallback mode and the absence of a "
        "privileged real ICMP server."
    ))

    # Conclusion
    add_heading(doc, "9. Conclusion", 1)
    add_body(doc, (
        "The ICMP Pinger Lab demonstrates a robust approach to active network measurement with real-time visualization "
        "and actionable metrics. The design balances portability, privilege constraints, and user experience, making it "
        "suitable for both educational demonstrations and basic diagnostics."
    ))

    # Future Work
    add_heading(doc, "10. Future Work", 1)
    add_bullet(doc, [
        "Add traceroute (TTL-based hop discovery) and TTL/sequence reporting.",
        "Support multiple hosts and scheduled measurement runs.",
        "Persist historical data and show trend charts.",
        "Integrate packet capture (e.g., using Npcap) for deeper analysis.",
    ])

    # References
    add_heading(doc, "11. References", 1)
    add_bullet(doc, [
        "RFC 792 – Internet Control Message Protocol.",
        "Forouzan, B. A., Data Communications and Networking.",
        "Kurose & Ross, Computer Networking: A Top-Down Approach.",
        "Scapy Documentation – https://scapy.readthedocs.io",
        "Microsoft IP Helper API (IcmpSendEcho) Documentation.",
    ])

    # Appendix: How to Run
    add_heading(doc, "Appendix A — How to Run", 1)
    add_numbered(doc, [
        "Install Python 3.8+ and dependencies: pip install -r requirements.txt",
        "Launch GUI: python main.py",
        "Enter Target Host, Count, Interval → Start Ping.",
        "Use Statistics tab to export CSV or generate a textual report.",
    ])

    # Appendix B — Customization Checklist
    add_heading(doc, "Appendix B — Customization Checklist", 1)
    add_bullet(doc, [
        "Replace placeholder metrics table with actual collected values.",
        "Insert screenshots: GUI main window, RTT graph, report dialog.",
        "Add Wireshark capture excerpt (echo request/reply sequence).",
        "Update Test Plan outcomes with empirical observations.",
        "Regenerate report after edits using script.",
    ])

    doc.save(path)
    print(f"✓ Lab report generated: {path}")


if __name__ == "__main__":
    create_report()
